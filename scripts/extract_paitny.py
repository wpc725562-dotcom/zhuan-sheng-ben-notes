"""提取 paitny/Upgrade 仓库资源 -> 资料/外部资源文字版/paitny/

零成本优先：文字层 PDF 直抽、docx 用 python-docx 读。
扫描件（3400 词汇 PDF）交给 ocr_forestdeer_batch.py 之外的单独 OCR 步骤。

源缓存：D:/专升本/_pdf_cache/paitny/（仓库外，不进 Git）
"""

from __future__ import annotations

import os
import re
import sys

import docx
import pymupdf

SRC = r"D:\专升本\_pdf_cache\paitny"
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "资料", "外部资源文字版", "paitny")

WATERMARK = re.compile(r"(叶学长|优课必过|zikao\.p1tao|prince-zip|网校出品|自考资料网)")


def clean_lines(text: str) -> str:
    kept = [l for l in text.splitlines() if not WATERMARK.search(l)]
    out = "\n".join(kept)
    return re.sub(r"\n{3,}", "\n\n", out).strip()


def from_pdf(path: str) -> str:
    doc = pymupdf.open(path)
    n = doc.page_count
    chunks = []
    for i, pg in enumerate(doc, 1):
        body = clean_lines(pg.get_text())
        if len(body) >= 30:
            chunks.append(f"\n<!-- 第 {i} 页 -->\n\n{body}\n")
    doc.close()
    return "\n".join(chunks), n


def from_docx(path: str) -> str:
    d = docx.Document(path)
    parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    # 表格也拉出来（题库常放表格里）
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_lines("\n\n".join(parts))


def main() -> int:
    os.makedirs(OUT, exist_ok=True)
    written = skipped = 0
    for f in sorted(os.listdir(SRC)):
        p = os.path.join(SRC, f)
        stem, ext = os.path.splitext(f)
        ext = ext.lower()
        dst = os.path.join(OUT, stem + ".md")
        try:
            if ext == ".pdf":
                body, n = from_pdf(p)
                if len(body) < 200:
                    print(f"SKIP {f}  <- 无文字层，需 OCR（{n}页）")
                    skipped += 1
                    continue
                head = f"# {stem}\n\n> 来源 paitny/Upgrade · 文字层直抽 · {n} 页\n\n"
            elif ext == ".docx":
                body = from_docx(p)
                if len(body) < 100:
                    print(f"SKIP {f}  <- 内容为空")
                    skipped += 1
                    continue
                head = f"# {stem}\n\n> 来源 paitny/Upgrade · python-docx 提取\n\n"
            else:
                print(f"SKIP {f}  <- 老 {ext} 格式，antiword 读不出（正文为嵌入图片）")
                skipped += 1
                continue
        except Exception as exc:  # noqa: BLE001
            print(f"BAD  {f}  <- {str(exc)[:70]}")
            skipped += 1
            continue
        with open(dst, "w", encoding="utf-8", newline="\n") as fh:
            fh.write(head + body + "\n")
        print(f"OK   {os.path.relpath(dst, REPO)}  ({len(body)} 字符)")
        written += 1
    print(f"\n写出 {written}, 跳过 {skipped}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
